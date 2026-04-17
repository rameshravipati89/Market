"""
DOCX text extraction using python-docx.
Pulls text from paragraphs and tables, returns a single plain-text string.
"""

import io

from docx import Document


def extract(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file given its raw bytes."""
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
